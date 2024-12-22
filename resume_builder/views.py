from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse, FileResponse
from .models import Resume
from .forms import ResumeForm
import json
from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from typing import List
import os
import docx
import PyPDF2
import tempfile
from django.conf import settings
import logging
logger = logging.getLogger(__name__)
from dotenv import load_dotenv
from django.urls import reverse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
from docx import Document
from docx.shared import Inches
from django.core.files.base import ContentFile
from google.cloud import storage
from django.conf import settings
import uuid

# Load environment variables from .env file

class OptimizedResume(BaseModel):
    optimized_content: str = Field(description="The ATS-optimized resume content")
    keyword_matches: List[str] = Field(description="List of important keywords matched from job description")
    improvement_suggestions: List[str] = Field(description="List of suggestions for improving the resume")
    ats_score: int = Field(description="ATS compatibility score out of 100")

class ResumeOptimizer:
    def __init__(self, api_key: str):
        self.llm = ChatOpenAI(
            temperature=0.1,
            model_name="gpt-4",
            openai_api_key=api_key
        )
        self.output_parser = PydanticOutputParser(pydantic_object=OptimizedResume)

    def optimize(self, resume_text: str, job_description: str) -> OptimizedResume:
        template = """
        You are an expert ATS (Applicant Tracking System) optimization specialist and professional resume writer.
        Analyze the provided resume and job description, then optimize the resume for ATS compatibility while 
        maintaining a professional tone.

        Resume:
        {resume_text}

        Job Description:
        {job_description}

        Please optimize the resume by:
        1. Identifying and incorporating relevant keywords from the job description
        2. Improving formatting for ATS readability
        3. Enhancing bullet points to better match job requirements
        4. Maintaining professional language and tone
        5. Ensuring proper section organization

        {format_instructions}
        """

        prompt = ChatPromptTemplate.from_template(template)
        
        messages = prompt.format_messages(
            resume_text=resume_text,
            job_description=job_description,
            format_instructions=self.output_parser.get_format_instructions()
        )

        response = self.llm.predict_messages(messages)
        return self.output_parser.parse(response.content)

def create_resume(request):
    # Hero section content
    hero_content = {
        'title': 'Resume Builder & Optimizer',
        'description': 'Create or optimize your ATS-friendly resume with our professional templates and AI-powered suggestions. Stand out to employers and increase your chances of getting hired.',
        'buttons': {
            'create': 'Create New Resume',
            'optimize': 'Optimize Existing Resume',
            'job_service': 'Job Application Service'
        },
        'credits_text': 'Credits per resume creation or optimization'
    }

    context = {
        'form': ResumeForm(),
        'hero_content': hero_content,
        'resume_services': {
            'builder_cost': 10  # or whatever the cost is
        }
    }

    return render(request, 'resume.html', context)

@login_required
def optimize_resume(request):
    if request.method == 'POST':
        try:
            # Debug logging
            logger.debug(f"Form data: {request.POST}")
            logger.debug(f"Files: {request.FILES}")

            resume_text = ""
            resume_file = request.FILES.get('resume_file')
            
            if resume_file:
                logger.info(f"Processing file: {resume_file.name}")
                # Get file extension
                file_extension = os.path.splitext(resume_file.name)[1].lower()
                
                # Create a temporary file to handle the upload
                with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                    for chunk in resume_file.chunks():
                        temp_file.write(chunk)
                
                try:
                    # Process different file types
                    if file_extension == '.pdf':
                        logger.info("Processing PDF file")
                        with open(temp_file.name, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            for page in pdf_reader.pages:
                                resume_text += page.extract_text()
                    elif file_extension in ['.doc', '.docx']:
                        logger.info("Processing Word document")
                        doc = docx.Document(temp_file.name)
                        resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    else:
                        raise ValueError(f"Unsupported file type: {file_extension}")
                except Exception as e:
                    logger.error(f"File processing error: {str(e)}")
                    raise
                finally:
                    # Clean up the temporary file
                    os.unlink(temp_file.name)
            else:
                logger.info("Using text input")
                resume_text = request.POST.get('resume_text', '').strip()

            if not resume_text:
                logger.error("No resume content provided")
                raise ValueError("Please provide either a resume file or paste resume text")

            job_description = request.POST.get('job_description', '').strip()
            if not job_description:
                logger.error("No job description provided")
                raise ValueError("Please provide a job description")

            logger.info("Starting resume optimization")
            optimizer = ResumeOptimizer(api_key=settings.OPENAI_API_KEY)
            optimization_result = optimizer.optimize(resume_text, job_description)
            logger.info("Resume optimization completed")

            # Store optimization results in session
            request.session['optimization_results'] = {
                'original_content': resume_text,
                'optimized_content': optimization_result.optimized_content,
                'job_description': job_description,
                'keyword_matches': optimization_result.keyword_matches,
                'improvement_suggestions': optimization_result.improvement_suggestions,
                'ats_score': optimization_result.ats_score
            }

            return JsonResponse({
                'success': True,
                'message': 'Resume optimized successfully',
                'redirect_url': reverse('resume_builder:download_resume')  # Use reverse() with named URL
            })

        except ValueError as ve:
            logger.warning(f"Validation error: {str(ve)}")
            return JsonResponse({
                'success': False,
                'error': str(ve)
            }, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}", exc_info=True)
            return JsonResponse({
                'success': False,
                'error': f"An error occurred: {str(e)}"
            }, status=400)
    
    # GET request
    hero_content = {
        'page_title': 'Optimize Your Resume',
        'page_description': 'Upload your resume and job description to create an ATS-optimized version tailored to the position.'
    }
    return render(request, 'resume_builder/optimize_resume.html', {'hero_content': hero_content})

@login_required
def download_resume(request):
    # Get resume data from session
    resume_data = request.session.get('resume_data')
    if not resume_data:
        return redirect('resume_builder:create_resume')
    
    # Available templates with descriptive IDs
    templates = [
        {'id': 'professional', 'name': 'Professional'},
        {'id': 'modern', 'name': 'Modern'},
        {'id': 'creative', 'name': 'Creative'},
    ]
    
    # Get active template (default to professional)
    active_template = request.session.get('active_template', 'professional')
    
    context = {
        'resume_data': resume_data,
        'templates': templates,
        'active_template': active_template,
        'selected_template': f'resume_templates/{active_template}.html'
    }
    
    return render(request, 'download_resume.html', context)

@login_required
def switch_template(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            template_id = data.get('template_id')
            
            # Update valid template IDs
            if template_id not in ['professional', 'modern', 'creative']:
                raise ValueError('Invalid template ID')
            
            # Store selected template in session
            request.session['active_template'] = template_id
            
            # Render the selected template with resume data
            resume_data = request.session.get('resume_data')
            return render(request, f'resume_templates/{template_id}.html', {'resume_data': resume_data})
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)

def resume_home(request):
    return render(request, 'resume_builder/home.html')

@login_required
def download_pdf(request):
    if request.method == 'POST':
        try:
            # Get the content from the request
            data = json.loads(request.body)
            content = data.get('content', '')

            # Create a file-like buffer to receive PDF data
            buffer = BytesIO()

            # Create the PDF object, using the buffer as its "file."
            p = canvas.Canvas(buffer, pagesize=letter)

            # Draw the content
            y = 750  # Starting y position
            for line in content.split('\n'):
                if line.strip():  # Only process non-empty lines
                    p.drawString(72, y, line.strip())
                    y -= 15  # Move down by 15 points
                    if y <= 50:  # Start a new page if we're near the bottom
                        p.showPage()
                        y = 750

            # Close the PDF object cleanly
            p.showPage()
            p.save()

            # Get the value of the BytesIO buffer and return it
            buffer.seek(0)
            return FileResponse(
                buffer, 
                as_attachment=True, 
                filename='optimized_resume.pdf',
                content_type='application/pdf'
            )

        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)

    return JsonResponse({
        'success': False,
        'error': 'Only POST requests are allowed'
    }, status=405)

@login_required
def download_word(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            content = data.get('content', '')
            save_to_cloud = data.get('save_to_cloud', False)

            # Create a new Word document
            doc = Document()
            doc.add_heading('Optimized Resume', 0)

            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            # Save to BytesIO instead of temporary file
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            if save_to_cloud:
                # Save to Google Cloud Storage
                try:
                    storage_client = storage.Client()
                    bucket = storage_client.bucket(settings.GCS_BUCKET_NAME)
                    
                    # Generate unique filename
                    filename = f"resumes/{uuid.uuid4()}.docx"
                    blob = bucket.blob(filename)
                    
                    # Upload from memory
                    blob.upload_from_file(
                        doc_io, 
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )

                    # Save reference to database
                    resume = Resume.objects.create(
                        user=request.user,
                        file_url=f"gs://{settings.GCS_BUCKET_NAME}/{filename}",
                        file_type='docx'
                    )

                    return JsonResponse({
                        'success': True,
                        'message': 'Resume saved to cloud',
                        'file_url': resume.file_url
                    })

                except Exception as e:
                    logger.error(f"Cloud storage error: {str(e)}")
                    return JsonResponse({
                        'success': False,
                        'error': 'Failed to save to cloud storage'
                    }, status=500)

            # For direct download
            response = HttpResponse(
                doc_io.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename=optimized_resume.docx'
            return response

        except Exception as e:
            logger.error(f"Word generation error: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)

    return JsonResponse({
        'success': False,
        'error': 'Only POST requests are allowed'
    }, status=405)



def save_resume(request):
    return HttpResponse(r'saved')

@login_required
def create_resume_submit(request):
    if request.method == 'POST':
        try:
            # Parse the JSON data from request
            data = json.loads(request.body)
            
            # Store resume data in session for template rendering
            request.session['resume_data'] = {
                'personal_info': {
                    'full_name': data.get('fullName'),
                    'title': data.get('title'),
                    'email': data.get('email'),
                    'phone': data.get('phone'),
                    'summary': data.get('summary')
                },
                'experience': data.get('experience', []),
                'education': data.get('education', []),
                'skills': {
                    'technical': data.get('technicalSkills', '').split(','),
                    'soft': data.get('softSkills', '').split(','),
                    'languages': data.get('languages', '').split(',')
                },
                'additional': {
                    'certifications': data.get('certifications'),
                    'projects': data.get('projects')
                }
            }

            return JsonResponse({
                'success': True,
                'message': 'Resume created successfully',
                'redirect_url': reverse('resume_builder:download_resume')
            })

        except json.JSONDecodeError as e:
            return JsonResponse({
                'success': False,
                'error': 'Invalid JSON data'
            }, status=400)
        except Exception as e:
            logger.error(f"Resume creation error: {str(e)}")
            return JsonResponse({
                'success': False,
                'error': 'Failed to create resume'
            }, status=500)

    return JsonResponse({
        'success': False,
        'error': 'Method not allowed'
    }, status=405)